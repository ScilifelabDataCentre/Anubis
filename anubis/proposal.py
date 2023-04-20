"""Proposal display and edit.

A proposal is created from an open call. It must be created by a user
account in the system. It may be transferred to another user. A user
may have at most one proposal in a call.

A proposal is defined by the proposal fields in the call.
"""

import io
import os.path

import docx
import flask
import htmldocx
import xlsxwriter

import anubis.call
import anubis.database
import anubis.decision
import anubis.grant
import anubis.review
import anubis.user
from anubis import constants
from anubis import utils
from anubis.saver import Saver, FieldSaverMixin, AccessSaverMixin


blueprint = flask.Blueprint("proposal", __name__)


@blueprint.route("/<pid>")
@utils.login_required
def display(pid):
    "Display the proposal."
    proposal = get_proposal(pid)
    if proposal is None:
        return utils.error("No such proposal.")
    if not allow_view(proposal):
        return utils.error("You are not allowed to view this proposal.")

    call = anubis.call.get_call(proposal["call"])
    am_submitter = (
        flask.g.current_user and flask.g.current_user["username"] == proposal["user"]
    )
    submitter_email = anubis.user.get_user(username=proposal["user"])["email"]
    access_emails = []
    for username in proposal.get("access_view", []):
        user = anubis.user.get_user(username=username)
        if user:
            access_emails.append(user["email"])
    # There may be accounts that have no email!
    access_emails = [e for e in access_emails if e]
    if submitter_email:
        all_emails = [submitter_email] + access_emails
    else:
        all_emails = access_emails
    email_lists = {
        "Proposal submitter": submitter_email,
        "Persons with access to this proposal": ", ".join(access_emails),
        "All involved persons": ", ".join(all_emails),
    }
    decision = anubis.decision.get_decision(proposal.get("decision"))
    # Only show decision in-line in proposal for non-admin or non-staff.
    allow_view_decision = (
        decision
        and decision.get("finalized")
        and not (flask.g.am_admin or flask.g.am_staff)
        and call["privileges"].get("allow_submitter_view_decision")
    )
    grant = anubis.grant.get_grant_proposal(proposal["identifier"])
    return flask.render_template(
        "proposal/display.html",
        proposal=proposal,
        call=call,
        decision=decision,
        grant=grant,
        n_reviews=anubis.database.get_count(
            "reviews", "proposal", proposal["identifier"]
        ),
        n_reviews_archived=anubis.database.get_count(
            "reviews", "proposal_archived", proposal["identifier"]
        ),
        email_lists=email_lists,
        allow_edit=allow_edit(proposal),
        allow_delete=allow_delete(proposal),
        allow_submit=allow_submit(proposal),
        allow_transfer=allow_transfer(proposal),
        am_submitter=am_submitter,
        am_reviewer=anubis.call.am_reviewer(call),
        my_review=anubis.review.get_reviewer_review(proposal, flask.g.current_user),
        allow_view_reviews=anubis.call.allow_view_reviews(call),
        allow_create_decision=anubis.decision.allow_create(proposal),
        allow_link_decision=anubis.decision.allow_link(decision),
        allow_view_decision=allow_view_decision,
        allow_create_grant=anubis.grant.allow_create(proposal),
        allow_link_grant=anubis.grant.allow_link(grant),
    )


@blueprint.route("/<pid>.docx")
@utils.login_required
def display_docx(pid):
    "Return a DOCX file containing the proposal information."
    proposal = get_proposal(pid)
    if proposal is None:
        return utils.error("No such proposal.")
    if not allow_view(proposal):
        return utils.error("You are not allowed to view this proposal.")

    content = get_proposal_docx(proposal).getvalue()
    response = flask.make_response(content)
    response.headers.set("Content-Type", constants.DOCX_MIMETYPE)
    response.headers.set(
        "Content-Disposition", "attachment", filename=f"{pid.replace(':','-')}.docx"
    )
    return response


@blueprint.route("/<pid>.xlsx")
@utils.login_required
def display_xlsx(pid):
    "Return an XLSX file containing the proposal information."
    proposal = get_proposal(pid)
    if proposal is None:
        return utils.error("No such proposal.")
    if not allow_view(proposal):
        return utils.error("You are not allowed to view this proposal.")

    content = get_proposal_xlsx(proposal).getvalue()
    response = flask.make_response(content)
    response.headers.set("Content-Type", constants.XLSX_MIMETYPE)
    response.headers.set(
        "Content-Disposition", "attachment", filename=f"{pid.replace(':','-')}.xlsx"
    )
    return response


@blueprint.route("/<pid>/edit", methods=["GET", "POST", "DELETE"])
@utils.login_required
def edit(pid):
    "Edit the proposal."
    proposal = get_proposal(pid)
    if proposal is None:
        return utils.error("No such proposal.")
    if not allow_edit(proposal):
        return utils.error("You are not allowed to edit this proposal.")

    call = anubis.call.get_call(proposal["call"])

    if utils.http_GET():
        return flask.render_template("proposal/edit.html", proposal=proposal, call=call)

    elif utils.http_POST():
        try:
            with ProposalSaver(proposal) as saver:
                saver["title"] = flask.request.form.get("_title") or None
                saver.set_fields_values(call["proposal"], form=flask.request.form)
        except ValueError as error:
            return utils.error(error)

        # If a repeat field was changed, then redisplay edit page.
        if saver.repeat_changed:
            return flask.redirect(
                flask.url_for("proposal.edit", pid=proposal["identifier"])
            )

        if flask.request.form.get("_save") == "submit":
            proposal = get_proposal(pid, refresh=True)  # Get up-to-date info.
            try:
                with ProposalSaver(proposal) as saver:
                    saver.set_submitted()  # Tests whether allowed or not.
            except ValueError as error:
                utils.flash_error(error)
            else:
                utils.flash_message("Proposal saved and submitted.")
                try:
                    send_email_submission(proposal)
                except ValueError:
                    utils.flash_warning("No separate confirmation email sent.")
                except KeyError:
                    utils.flash_error(
                        "Could not send confirmation email; misconfiguration in the Anubis setup."
                    )

        elif allow_submit(proposal) and not proposal.get("submitted"):
            utils.flash_warning(
                "Proposal was saved but not submitted."
                " You must explicitly submit it!"
            )
        return flask.redirect(
            flask.url_for("proposal.display", pid=proposal["identifier"])
        )

    elif utils.http_DELETE():
        if not allow_delete(proposal):
            return utils.error("You are not allowed to delete this proposal.")
        decision = anubis.decision.get_decision(proposal.get("decision"))
        if decision:
            anubis.database.delete(decision)
        reviews = anubis.database.get_docs(
            "reviews", "proposal", proposal["identifier"]
        )
        for review in reviews:
            anubis.database.delete(review)
        anubis.database.delete(proposal)
        utils.flash_message(f"Deleted proposal {pid}.")
        if flask.g.am_admin or flask.g.am_staff:
            url = flask.url_for("proposals.call", cid=call["identifier"])
        else:
            url = flask.url_for("proposals.user", username=proposal["user"])
        return flask.redirect(url)


@blueprint.route("/<pid>/transfer", methods=["GET", "POST"])
@utils.login_required
def transfer(pid):
    "Transfer ownership of he proposal."
    proposal = get_proposal(pid)
    if proposal is None:
        return utils.error("No such proposal.")
    if not allow_transfer(proposal):
        return utils.error(
            "You are not allowed to transfer ownership of this proposal."
        )

    if utils.http_GET():
        return flask.render_template("proposal/transfer.html", proposal=proposal)

    elif utils.http_POST():
        try:
            with ProposalSaver(proposal) as saver:
                value = flask.request.form.get("user")
                if value:
                    user = anubis.user.get_user(username=value, email=value)
                    if user:
                        saver.set_user(user)
                    else:
                        raise ValueError("No such user.")
        except ValueError as error:
            return utils.error(error)
        return flask.redirect(
            flask.url_for("proposal.display", pid=proposal["identifier"])
        )


@blueprint.route("/<pid>/submit", methods=["POST"])
@utils.login_required
def submit(pid):
    "Submit the proposal."
    proposal = get_proposal(pid)
    if proposal is None:
        return utils.error("No such proposal.")

    if utils.http_POST():
        try:
            with ProposalSaver(proposal) as saver:
                saver.set_submitted()  # Tests whether allowed or not.
        except ValueError as error:
            utils.flash_error(error)
        else:
            utils.flash_message("Proposal was submitted.")
            try:
                send_email_submission(proposal)
            except ValueError:
                utils.flash_warning("No separate confirmation email sent.")
            except KeyError:
                utils.flash_error(
                    "Could not send confirmation email; misconfiguration in the Anubis setup."
                )
        return flask.redirect(flask.url_for("proposal.display", pid=pid))


def send_email_submission(proposal):
    """Send an email to the owner of the proposal confirming the submission.
    Raise ValueError if email server not configured.
    Raise KeyError if email could not be sent; server misconfigured.
    """
    user = anubis.user.get_user(username=proposal["user"])
    if not (user and user["email"]):
        return
    site = flask.current_app.config["SITE_NAME"]
    title = f"Proposal {proposal['identifier']} has been submitted in {site}"
    url = flask.url_for("proposal.display", pid=proposal["identifier"], _external=True)
    text = (
        "Your proposal\n\n"
        f"  {proposal['identifier']} {proposal['title']}\n\n"
        f"has been submitted in the {site} system.\n\n"
        f"View it at {url}\n\n"
        "/The Anubis system"
    )
    utils.send_email(user["email"], title, text)


@blueprint.route("/<pid>/unsubmit", methods=["POST"])
@utils.login_required
def unsubmit(pid):
    "Unsubmit the proposal."
    proposal = get_proposal(pid)
    if proposal is None:
        return utils.error("No such proposal.")

    if utils.http_POST():
        try:
            with ProposalSaver(proposal) as saver:
                saver.set_unsubmitted()  # Tests whether allowed or not.
        except ValueError as error:
            utils.flash_error(error)
        else:
            utils.flash_warning("Proposal was unsubmitted.")
        return flask.redirect(flask.url_for("proposal.display", pid=pid))


@blueprint.route("/<pid>/access", methods=["GET", "POST", "DELETE"])
@utils.login_required
def change_access(pid):
    "Edit the access privileges for the proposal record."
    proposal = get_proposal(pid)
    if proposal is None:
        return utils.error("No such proposal.")
    if not allow_edit(proposal):
        return utils.error("You are not allowed to edit this proposal.")
    call = anubis.call.get_call(proposal["call"])

    if utils.http_GET():
        users_edit = sorted(proposal.get("access_edit", []))
        users_view = sorted(set(proposal.get("access_view", [])).difference(users_edit))
        return flask.render_template(
            "change_access.html",
            title=f"Proposal {proposal['identifier']}",
            url=flask.url_for("proposal.change_access", pid=proposal["identifier"]),
            users_view=users_view,
            users_edit=users_edit,
            back_url=flask.url_for("proposal.display", pid=proposal["identifier"]),
        )

    elif utils.http_POST():
        try:
            with ProposalSaver(doc=proposal) as saver:
                saver.set_access(form=flask.request.form)
        except ValueError as error:
            utils.flash_error(error)
        return flask.redirect(
            flask.url_for("proposal.change_access", pid=proposal["identifier"])
        )

    elif utils.http_DELETE():
        try:
            with ProposalSaver(doc=proposal) as saver:
                saver.remove_access(form=flask.request.form)
        except ValueError as error:
            utils.flash_error(error)
        return flask.redirect(
            flask.url_for("proposal.change_access", pid=proposal["identifier"])
        )


@blueprint.route("/<pid>/document/<fid>")
@utils.login_required
def document(pid, fid):
    "Download the proposal document (attachment file) for the given field id."
    proposal = get_proposal(pid)
    if proposal is None:
        return utils.error("No such proposal.")
    if not allow_view(proposal):
        return utils.error("You are not allowed to read this proposal.")
    try:
        doc = get_document(proposal, fid)
    except KeyError:
        return utils.error(
            "No such document in the proposal.",
            flask.url_for("proposal.display", pid=pid),
        )

    response = flask.make_response(doc["content"])
    response.headers.set("Content-Type", doc["content_type"])
    response.headers.set("Content-Disposition", "attachment", filename=doc["filename"])
    return response


def get_document(proposal, fid):
    "Return a dictionary containing the document in the field of the proposal."
    documentname = proposal["values"][fid]
    # This may generate a KeyError, which is correct.
    stub = proposal["_attachments"][documentname]
    # Colon ':' is a problematic character in filenames.
    # Replace it by dash '-' which used as general glue character here.
    pid = proposal["identifier"].replace(":", "-")
    ext = os.path.splitext(documentname)[1]
    outfile = flask.g.db.get_attachment(proposal, documentname)
    return dict(
        filename=f"{pid}-{fid}{ext}",
        content=outfile.read(),
        content_type=stub["content_type"],
    )


@blueprint.route("/<pid>/logs")
@utils.login_required
def logs(pid):
    "Display the log records of the given proposal."
    proposal = get_proposal(pid)
    if proposal is None:
        return utils.error("No such proposal.")
    if not allow_view(proposal):
        return utils.error("You are not allowed to read this proposal.")

    return flask.render_template(
        "logs.html",
        title=f"Proposal {proposal['identifier']}",
        back_url=flask.url_for("proposal.display", pid=proposal["identifier"]),
        logs=anubis.database.get_logs(proposal["_id"]),
    )


class ProposalSaver(AccessSaverMixin, FieldSaverMixin, Saver):
    "Proposal document saver context."

    DOCTYPE = constants.PROPOSAL

    def __init__(self, doc=None, call=None, user=None):
        if doc:
            super().__init__(doc=doc)
        elif call and user:
            super().__init__(doc=None)
            self.set_call(call)
            self.set_user(user)
        else:
            raise ValueError("doc or call+user must be specified")

    def initialize(self):
        self.doc["values"] = {}
        self.doc["errors"] = {}
        self.doc["access_view"] = []
        self.doc["access_edit"] = []

    def set_user(self, user):
        "Set the user (owner) for the proposal; must be called when creating."
        if get_call_user_proposal(self.doc["call"], user["username"]):
            raise ValueError("User already has a proposal in the call.")
        self.doc["user"] = user["username"]

    def set_call(self, call):
        "Set the call for the proposal; must be called when creating."
        if self.doc.get("call"):
            raise ValueError("call has already been set")
        self.doc["call"] = call["identifier"]
        counter = call.get("counter")
        if counter is None:
            counter = 1
        else:
            counter += 1
        with anubis.call.CallSaver(call):
            call["counter"] = counter
        self.doc["identifier"] = f"{call['identifier']}:{counter:03d}"
        self.set_fields_values(call["proposal"])

    def set_submitted(self):
        if not allow_submit(self.doc):
            raise ValueError(
                "Submit cannot be done; proposal is incomplete," " or call is closed."
            )
        self.doc["submitted"] = utils.get_now()

    def set_unsubmitted(self):
        if not allow_submit(self.doc):
            raise ValueError("Unsubmit cannot be done; call is closed.")
        self.doc.pop("submitted", None)


def get_proposal(pid, refresh=False):
    """Return the proposal with the given identifier.
    Return None if not found.
    """
    key = f"proposal {pid}"
    try:
        if refresh:
            raise KeyError
        return utils.cache_get(key)
    except KeyError:
        docs = [
            r.doc
            for r in flask.g.db.view(
                "proposals", "identifier", key=pid, include_docs=True
            )
        ]
        if len(docs) == 1:
            return utils.cache_put(key, docs[0])
        else:
            return None


def get_call_user_proposal(cid, username):
    "Return the proposal owned by the user in the call."
    result = [
        r.doc
        for r in flask.g.db.view(
            "proposals", "call_user", key=[cid, username], include_docs=True
        )
    ]
    if len(result) == 1:
        return result[0]
    else:
        return None


def get_proposal_docx(proposal):
    "Return the proposal as a io.BytesIO instance containing the DOCX file."
    call = anubis.call.get_call(proposal["call"])
    am_submitter = (
        flask.g.current_user and flask.g.current_user["username"] == proposal["user"]
    )
    submitter = anubis.user.get_user(username=proposal["user"])
    doc = docx.Document()
    doc.add_heading(f"Proposal {proposal['identifier']}", 0)
    doc.add_heading(proposal["title"], 1)
    para = doc.add_paragraph()
    para.paragraph_format.space_before = docx.shared.Pt(20)
    para.add_run("Submitter: ").bold = True
    para.add_run(anubis.user.get_fullname(submitter))
    para.add_run(f" (Anubis user name: {submitter['username']})")
    para = doc.add_paragraph()
    para.add_run("Affiliation: ").bold = True
    para.add_run(submitter.get("affiliation") or "-")
    para = doc.add_paragraph()
    para.add_run("Modified: ").bold = True
    para.add_run(proposal["modified"])
    para = doc.add_paragraph()
    para.add_run("Call: ").bold = True
    para.add_run(f"{call['identifier']}: {call['title']}")
    para = doc.add_paragraph()
    para.add_run("Proposal URL: ").bold = True
    para.add_run(
        flask.url_for("proposal.display", pid=proposal["identifier"], _external=True)
    )
    for field in call["proposal"]:
        doc.add_heading(field["title"] or field["identifier"].capitalize(), 2)
        value = proposal["values"].get(field["identifier"])
        if value is None:
            doc.add_paragraph("-")
        elif field["type"] in (
            constants.LINE,
            constants.EMAIL,
        ):
            doc.add_paragraph(value)
        elif field["type"] == constants.BOOLEAN:
            if value is None:
                value = "-"
            else:
                value = value and "Yes" or "No"
            doc.add_paragraph(value)
        elif field["type"] == constants.SELECT:
            if isinstance(value, list):
                doc.add_paragraph("; ".join(value))
            else:
                doc.add_paragraph(value)
        elif field["type"] in (
            constants.INTEGER,
            constants.FLOAT,
            constants.SCORE,
            constants.RANK,
        ):
            doc.add_paragraph(str(value))
        elif field["type"] == constants.TEXT:
            converter = htmldocx.HtmlToDocx()
            converter.add_html_to_document(utils.markdown2html(value), doc)
        elif field["type"] == constants.DOCUMENT:
            para = doc.add_paragraph()
            para.add_run("Document: ").bold = True
            documentname = proposal["values"][field["identifier"]]
            pid = proposal["identifier"].replace(":", "-")
            ext = os.path.splitext(documentname)[1]
            para.add_run(f"{pid}-{field['identifier']}{ext}")
            para.add_run(f' (originally: "{documentname}")')
        else:
            pass  # Ignore unimplemented field types.
    result = io.BytesIO()
    doc.save(result)
    return result


def get_proposal_xlsx(proposal):
    "Return the proposal as a io.BytesIO instance containing the XLSX file."
    call = anubis.call.get_call(proposal["call"])
    am_submitter = (
        flask.g.current_user and flask.g.current_user["username"] == proposal["user"]
    )
    submitter = anubis.user.get_user(username=proposal["user"])
    result = io.BytesIO()
    wb = xlsxwriter.Workbook(result, {"in_memory": True})
    formats = utils.create_xlsx_formats(wb)
    # Hard str(len) limit for worksheet title.
    ws = wb.add_worksheet(f"Proposal {proposal['identifier'].replace(':','-')}"[:31])
    ws.set_column(0, 0, 20, formats["head"])
    ws.set_column(1, 1, 80, formats["normal"])
    ws.set_column(2, 2, 60, formats["normal"])
    nrow = 0
    row = ["Proposal", "", proposal["title"]]
    ws.write_row(nrow, 0, row)
    ws.write_url(
        nrow,
        1,
        flask.url_for("proposal.display", pid=proposal["identifier"], _external=True),
        string=proposal["identifier"],
    )
    nrow += 1
    row = [
        "Submitter",
        anubis.user.get_fullname(submitter),
        f"{submitter.get('affiliation') or '-'}",
    ]
    ws.write_row(nrow, 0, row)
    nrow += 1
    row = ["Modified", proposal["modified"]]
    ws.write_row(nrow, 0, row)
    nrow += 1
    row = ["Call", "", call["title"]]
    ws.write_url(
        nrow,
        1,
        flask.url_for("call.display", cid=call["identifier"], _external=True),
        string=call["identifier"],
    )
    ws.write_row(nrow, 0, row)
    nrow += 2
    for field in call["proposal"]:
        ws.write_string(nrow, 0, field["title"] or field["identifier"].capitalize())
        value = proposal["values"].get(field["identifier"])
        # Ugly, but necessary...
        if value is not None and field["type"] == constants.DOCUMENT:
            value = flask.url_for(
                "proposal.document",
                pid=proposal["identifier"],
                fid=field["identifier"],
                _external=True,
            )
        utils.write_xlsx_field(ws, nrow, 1, value, field["type"], formats)
        nrow += 1
    wb.close()
    return result


def allow_create(call):
    """A logged-in user may create a proposal in a call.
    Admin, staff, call owner and user with access to the call may always
    create a proposal. A reviewer in the call may not.
    Others may create a proposal only if the call is open and not closed.
    """
    if not flask.g.current_user:
        return False
    if flask.g.am_admin:
        return True
    if flask.g.am_staff:
        return True
    if anubis.call.am_owner(call):
        return True
    if anubis.call.am_reviewer(call):
        return False
    if anubis.call.allow_view(call):
        return anubis.call.is_open(call)
    return False


def allow_view(proposal):
    """The admin, staff and call owner may view a proposal.
    The user (owner) of the proposal may view it.
    A user set to have view access may view it.
    The reviewers may view it if it is submitted.
    """
    if not flask.g.current_user:
        return False
    if flask.g.am_admin:
        return True
    if flask.g.am_staff:
        return True
    if flask.g.current_user["username"] == proposal["user"]:
        return True
    call = anubis.call.get_call(proposal["call"])
    if anubis.call.am_owner(call):
        return True
    if flask.g.current_user["username"] in proposal.get("access_view", []):
        return True
    if anubis.call.am_reviewer(call):
        return bool(proposal.get("submitted"))
    return False


def allow_edit(proposal):
    """The admin and call owner may edit the proposal.
    The user may edit if not submitted.
    A user set to have edit access may edit it if not submitted.
    """
    if not flask.g.current_user:
        return False
    if flask.g.am_admin:
        return True
    call = anubis.call.get_call(proposal["call"])
    if anubis.call.am_owner(call):
        return True
    if proposal.get("submitted"):
        return False
    if flask.g.current_user["username"] == proposal["user"]:
        return True
    if flask.g.current_user["username"] in proposal.get("access_edit", []):
        return True
    return False


def allow_delete(proposal):
    """The admin and call owner may delete the proposal.
    The user may delete if not submitted.
    """
    if not flask.g.current_user:
        return False
    if flask.g.am_admin:
        return True
    call = anubis.call.get_call(proposal["call"])
    if anubis.call.am_owner(call):
        return True
    if proposal.get("submitted"):
        return False
    if flask.g.current_user["username"] == proposal["user"]:
        return True
    return False


def allow_submit(proposal):
    """Only if there are no errors.
    The admin, staff and call owner may submit/unsubmit the proposal.
    The user may submit/unsubmit the proposal if the call is open.
    """
    if not flask.g.current_user:
        return False
    if proposal["errors"]:
        return False
    if flask.g.am_admin:
        return True
    if flask.g.am_staff:
        return True
    call = anubis.call.get_call(proposal["call"])
    if anubis.call.am_owner(call):
        return True
    if flask.g.current_user["username"] != proposal["user"]:
        return False
    if not anubis.call.is_open(call):
        return False
    return True


def allow_transfer(proposal):
    "The admin staff and call owner may transfer ownership of a proposal."
    if not flask.g.current_user:
        return False
    if flask.g.am_admin:
        return True
    if flask.g.am_staff:
        return True
    call = anubis.call.get_call(proposal["call"])
    if anubis.call.am_owner(call):
        return True
    return False
